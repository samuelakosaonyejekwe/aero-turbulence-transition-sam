"""
gen_equations.py
Build model.equations.docx : every governing equation of the UTSS solver,
written as NATIVE Word equations (LaTeX -> MathML -> OMML).  No images.
Also (re)writes 07_equations/equations_index.csv holding the LaTeX source,
and removes any legacy equation PNGs.

Author: Akosa Samuel Onyejekwe, 2026.
"""
import os, glob
import pandas as pd
import latex2mathml.converter as L
import mathml2omml
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

INK   = RGBColor(0x1d,0x2f,0x45)   # navy ink  (never black)
BLUE  = RGBColor(0x1b,0x49,0x65)
ROSE  = RGBColor(0x9b,0x2d,0x55)
SOFT  = RGBColor(0x3c,0x50,0x66)
EQD   = "07_equations"; os.makedirs(EQD, exist_ok=True)
MATHNS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# (key, section, descriptive title, LaTeX)
EQS = [
 # ---- 4.1  Inviscid panel flow ---------------------------------------
 ("E01","4.1  Inviscid panel flow","Pressure coefficient (inviscid)",
   r"C_p = 1 - \left(\frac{V}{U_\infty}\right)^{2}"),
 ("E04","4.1  Inviscid panel flow","Karman-Tsien compressibility correction",
   r"C_p = \frac{C_{p,0}}{\beta + \left(\frac{M_\infty^{2}}{1+\beta}\right)\frac{C_{p,0}}{2}}, \quad \beta = \sqrt{1-M_\infty^{2}}"),
 ("E02","4.1  Inviscid panel flow","Vortex-panel surface velocity (Kuethe & Chow)",
   r"\frac{V_{t_i}}{U_\infty} = \cos(\theta_i-\alpha) + \sum_{j=1}^{N} C_{t,ij}\,\gamma_j"),
 ("E03","4.1  Inviscid panel flow","Kutta condition (sharp trailing edge)",
   r"\gamma_{1} + \gamma_{N+1} = 0"),
 ("E05","4.2  Laminar boundary layer","Thwaites momentum thickness (laminar)",
   r"\theta^{2} = \frac{0.45\,\nu}{U_e^{6}} \int_{0}^{x} U_e^{5}\, dx'"),
 ("E06","4.2  Laminar boundary layer","Thwaites pressure-gradient and momentum-thickness Reynolds number",
   r"\lambda = \frac{\theta^{2}}{\nu}\frac{dU_e}{dx}, \quad Re_\theta = \frac{U_e\,\theta}{\nu}"),
 ("E06b","4.2  Laminar boundary layer","Kinetic-energy integral (two-equation laminar march)",
   r"\theta\frac{dH^{*}}{dx} = 2C_D - H^{*}\frac{C_f}{2} - H^{*}(1-H)\frac{\theta}{U_e}\frac{dU_e}{dx}"),
 ("E06c","4.2  Laminar boundary layer","Laminar closure functions (from the Falkner-Skan family)",
   r"H^{*}(H) = \frac{\theta^{*}}{\theta}, \quad l(H) = Re_\theta\frac{C_f}{2} = \theta_\eta f''(0), "
   r"\quad d(H) = Re_\theta C_D = \theta_\eta \int \left(f''\right)^{2} d\eta"),
 ("E07","4.2  Laminar boundary layer","Von Karman momentum-integral equation",
   r"\frac{d\theta}{dx} + (2+H)\frac{\theta}{U_e}\frac{dU_e}{dx} = \frac{C_f}{2}"),
 # ---- 4.3  Unified transition kernel (novel) -------------------------
 ("E08","4.3  Unified transition kernel (novel)","Abu-Ghannam & Shaw bypass onset",
   r"Re_{\theta t} = 163 + \exp\left[F(\lambda_\theta) - \frac{F(\lambda_\theta)\,Tu}{6.91}\right]"),
 ("E09","4.3  Unified transition kernel (novel)","AGS pressure-gradient function",
   r"F(\lambda_\theta) = \begin{cases} 6.91+12.75\lambda_\theta+63.64\lambda_\theta^{2}, & \lambda_\theta \le 0 \\ "
   r"6.91+2.48\lambda_\theta-12.27\lambda_\theta^{2}, & \lambda_\theta > 0 \end{cases}"),
 ("E10","4.3  Unified transition kernel (novel)","Natural / Tollmien-Schlichting onset (e^N, envelope over frequency)",
   r"N(x) = \max_{\omega} \int_{x_0(\omega)}^{x} \frac{\sigma(H, Re_\theta, \omega\theta/U_e)}{\theta}\,dx' \ge N_{crit}"),
 ("E10b","4.3  Unified transition kernel (novel)","Orr-Sommerfeld operator (tabulated amplification rates)",
   r"\left[(U-c)(D^2-\alpha^2) - U''\right]\hat{v} = \frac{1}{i\alpha Re_\theta}(D^2-\alpha^2)^2\hat{v}"),
 ("E10c","4.3  Unified transition kernel (novel)","Gaster transformation (temporal to spatial growth rate)",
   r"\sigma = -\alpha_i\theta = \frac{\omega_i}{c_g}, \quad c_g = \frac{\partial\omega_r}{\partial\alpha_r}"),
 ("E11","4.3  Unified transition kernel (novel)","Cross-flow criterion (swept wing, C1 on Re_theta2)",
   r"Re_{\theta 2} = k_{cf}\,Re_\theta\,\sin\Lambda\,\cos\Lambda \ \ge\ C_1"),
 ("E11b","4.3  Unified transition kernel (novel)","Cross-flow amplification integral (the closure actually used)",
   r"N_{cf} = \int_{x_{c1}}^{x} \frac{\sigma(H_{rev})}{\theta}\,dx' \ \ge\ N_{crit}"),
 ("E12","4.3  Unified transition kernel (novel)","Separation bubble: dead-air march (momentum and kinetic energy, no wall shear)",
   r"\frac{d\theta}{dx} = -(2+H)\frac{\theta}{U_e}\frac{dU_e}{dx}, \quad "
   r"\theta\frac{dH^{*}}{dx} = 2C_D - H^{*}(1-H)\frac{\theta}{U_e}\frac{dU_e}{dx}, \quad C_f = 0"),
 ("E12b","4.3  Unified transition kernel (novel)","Separation bubble: reattachment condition",
   r"N_{bub} = \int_{x_s}^{x_r} \frac{\sigma(H_{rev}, Re_\theta)}{\theta}\,dx' = N_{crit}, \quad "
   r"\sigma(H_{rev}) \approx 0.0435"),
 ("E13","4.3  Unified transition kernel (novel)","Unified minimum-onset transition kernel",
   r"Re_{\theta t}^{*} = \min\left(a_{TS}Re_{\theta t}^{TS},\, a_{BP}Re_{\theta t}^{BP},\, "
   r"a_{SEP}Re_{\theta t}^{SEP},\, a_{CF}Re_{\theta t}^{CF}\right)"),
 ("E14","4.3  Unified transition kernel (novel)","Transition trigger",
   r"Re_\theta(x_t) \ge Re_{\theta t}^{*}"),
 # ---- 4.4  Transitional & turbulent boundary layer -------------------
 ("E15","4.4  Transitional & turbulent boundary layer","Narasimha universal intermittency",
   r"\gamma(x) = 1 - \exp\left[-0.412\,\xi^{2}\right], \quad \xi = \frac{x-x_t}{\lambda_{tr}}"),
 ("E16","4.4  Transitional & turbulent boundary layer","Transition-length scale",
   r"\lambda_{tr} = \frac{\nu}{U_e} C_{len}\, Re_{x,t}^{0.75}, \quad C_{len} = 9"),
 ("E17","4.4  Transitional & turbulent boundary layer","Intermittency-weighted property blend",
   r"\phi = (1-\gamma)\phi_{lam} + \gamma\,\phi_{turb}"),
 ("E18","4.4  Transitional & turbulent boundary layer","Head entrainment (turbulent)",
   r"\frac{d(\theta H_1)}{dx} = C_E - \frac{\theta H_1}{U_e}\frac{dU_e}{dx}, \quad C_E = 0.0306(H_1-3)^{-0.6169}"),
 ("E19","4.4  Transitional & turbulent boundary layer","Ludwieg-Tillmann skin-friction law",
   r"C_f = 0.246 \times 10^{-0.678 H}\, Re_\theta^{-0.268}"),
 # ---- 4.5  Drag, compressible temperature & reference ----------------
 ("E20","4.5  Drag, temperature & reference scales","Squire-Young profile drag",
   r"C_d = 2\frac{\theta_{TE}}{c}\left(\frac{U_{e,TE}}{U_\infty}\right)^{(H_{TE}+5)/2}"),
 ("E21","4.5  Drag, temperature & reference scales","Crocco-Busemann temperature profile",
   r"\frac{T}{T_e} = 1 + r\frac{\gamma-1}{2}M_e^{2}\left[1-\left(\frac{u}{U_e}\right)^{2}\right]"),
 ("E22","4.5  Drag, temperature & reference scales","Recovery (adiabatic-wall) temperature",
   r"T_r = T_e\left(1 + r\frac{\gamma-1}{2}M_e^{2}\right), \quad r \approx Pr^{1/3}"),
 ("E22b","4.5  Drag, temperature & reference scales","Eckert reference temperature (compressible closures)",
   r"\frac{T_{ref}}{T_e} = 1 + 0.032 M_e^{2} + 0.58\left(\frac{T_w}{T_e}-1\right), \quad "
   r"\frac{\nu_{ref}}{\nu_e} = \left(\frac{T_{ref}}{T_e}\right)^{1+\omega}"),
 ("E23","4.5  Drag, temperature & reference scales","Reynolds number (mean aerodynamic chord)",
   r"Re_{MAC} = \frac{\rho_\infty U_\infty \overline{c}}{\mu_\infty} = \frac{U_\infty \overline{c}}{\nu_\infty}"),
 ("E24","4.5  Drag, temperature & reference scales","Mean aerodynamic chord",
   r"\overline{c} = \frac{2}{3} c_{root}\frac{1+\lambda+\lambda^{2}}{1+\lambda}"),
]


def latex_to_omml_element(latex):
    """LaTeX -> MathML -> OMML -> lxml element ready to append to a paragraph."""
    mml  = L.convert(latex)
    omml = mathml2omml.convert(mml)
    if "xmlns:m" not in omml:
        omml = omml.replace("<m:oMath",
                            '<m:oMath xmlns:m="%s"' % MATHNS, 1)
    return parse_xml(omml)


def _set_font(run, size, color, bold=False, italic=False):
    run.font.name = "Cambria"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def build_doc():
    doc = Document()
    # base style: navy, never black
    base = doc.styles["Normal"]
    base.font.name = "Cambria"; base.font.size = Pt(11)
    base.font.color.rgb = INK

    # ---- title block --------------------------------------------------
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("UTSS — Governing Equations of the Model")
    _set_font(r, 20, BLUE, bold=True)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Universal Transition & Skin-friction Solver")
    _set_font(r, 12, SOFT, italic=True)
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run("Author: Akosa Samuel Onyejekwe   ·   Document UTSS-CASE-2026")
    _set_font(r, 10.5, SOFT)
    lead = doc.add_paragraph(); lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = lead.add_run("All equations below are rendered as native, editable Word equations.")
    _set_font(r, 10, SOFT, italic=True)
    doc.add_paragraph()

    rows = []
    last_section = None
    for key, section, title, latex in EQS:
        if section != last_section:
            h = doc.add_paragraph()
            rh = h.add_run(section)
            _set_font(rh, 14, BLUE, bold=True)
            # thin rule under the section heading
            p_pr = h._p.get_or_add_pPr()
            pbdr = parse_xml(
                '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:bottom w:val="single" w:sz="6" w:space="2" w:color="9BB0C4"/></w:pBdr>')
            p_pr.append(pbdr)
            last_section = section

        # caption line:  (E0x)  Title
        cap = doc.add_paragraph()
        rc = cap.add_run("(%s)  " % key); _set_font(rc, 10.5, ROSE, bold=True)
        rt = cap.add_run(title); _set_font(rt, 10.5, SOFT, italic=True)
        cap.paragraph_format.space_after = Pt(2)

        # the equation itself, centred
        eqp = doc.add_paragraph(); eqp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eqp.paragraph_format.space_after = Pt(10)
        try:
            el = latex_to_omml_element(latex)
            eqp._p.append(el)
        except Exception as e:
            fr = eqp.add_run("[equation %s] %s" % (key, latex))
            _set_font(fr, 11, INK)
            print("  ! fallback on", key, "->", e)

        rows.append((key, section.strip(), title, latex))

    out = "model.equations.docx"
    doc.save(out)
    print("wrote", out, "with", len(rows), "native equations")

    # ---- index csv (LaTeX source, not images) -------------------------
    pd.DataFrame(rows, columns=["key", "section", "equation", "latex"]).to_csv(
        f"{EQD}/equations_index.csv", index=False)
    print("wrote", f"{EQD}/equations_index.csv")

    # ---- remove legacy equation PNGs ----------------------------------
    removed = 0
    for png in glob.glob(f"{EQD}/eq_*.png"):
        os.remove(png); removed += 1
    if removed:
        print("removed", removed, "legacy equation images from", EQD)


if __name__ == "__main__":
    build_doc()
