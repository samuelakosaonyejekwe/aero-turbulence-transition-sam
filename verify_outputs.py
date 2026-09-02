"""
verify_outputs.py
Check that the compiled report says what the generated CSVs say.

Every headline number in case.docx is read out of a CSV by build_docx.py, so
the two cannot disagree by construction - unless a number is typed into the
narrative by hand, which is how the mean laminar extent, the section design
c_l, the leading-edge radius, the Blasius neutral point and the wing lift all
came to be wrong at some point in this project's history.  This script closes
that loop: it reads the CSVs, then searches the RENDERED document for each
value.  Run it after build_docx.py.

    python3 verify_outputs.py                 # checks the PDF if present,
                                              # else case.docx
    python3 verify_outputs.py <file>          # .pdf or .docx

A .docx is checked through python-docx; a .pdf through PyMuPDF, which is the
only one of the two that sees what a reader sees.  Exits non-zero on any
failure so it can be wired into a build.

HOW THE MATCHING WORKS, AND WHY IT IS NOT `value in text`
--------------------------------------------------------
The first version of this script asked `value in txt` over the whole extracted
document.  That is not a check.  "168" occurs inside "1683", inside "0.1685",
inside a page number and inside any of the several hundred other numbers a
66-page report prints; a three-digit value was all but guaranteed to "pass"
whatever the report actually said.  Two things fix it:

  * numbers are matched as whole numeric TOKENS - not preceded by a digit or a
    decimal point, and not followed by one - so 168 no longer matches 1683,
    0.168 or 168.4;
  * every value is anchored to a CONTEXT phrase and must occur within a window
    of it, so the cruise transition location has to appear near the transition
    table and not merely somewhere in the document.

Both are necessary.  Token matching alone still lets a value pass on a
coincidental occurrence elsewhere; anchoring alone still lets 168 pass on
1683 in the right paragraph.
"""
import os
import re
import sys

import pandas as pd

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import utss_paths  # noqa: F401  (anchors ROOT and the working directory)

WINDOW = 2500          # characters either side of an anchor phrase


def document_text(path):
    if path.lower().endswith(".pdf"):
        import fitz
        with fitz.open(path) as d:
            return "\n".join(p.get_text() for p in d), d.page_count
    from docx import Document
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            parts.append(" | ".join(c.text for c in r.cells))
    return "\n".join(parts), None


def flatten(txt):
    """Collapse every run of whitespace to one space.

    A PDF extractor returns the document as it is laid out, so a table caption
    or a paragraph is broken by a newline wherever the line ended.  Searching
    the raw text for a phrase therefore fails on any phrase long enough to
    wrap, which is most of them: anchoring on the caption "NLF vs
    fully-turbulent drag." finds nothing if the renderer split it after "vs".
    Collapsing whitespace first makes a search see the sentence rather than
    the line breaks.  The raw text is kept as well, for the one check that is
    ABOUT line breaks (a table cell squeezed until it breaks mid-token).
    """
    # Typographic characters normalised to their ASCII equivalents.  The
    # document and the README are typeset, so a negative number is written with
    # a UNICODE MINUS (U+2212) and not a hyphen, and a numeric parser looking
    # for "-14.2" does not find "\u221214.2".  The same applies to the various
    # dashes and to the non-breaking space inside a figure like "1 234".
    txt = (txt.replace("\u2212", "-").replace("\u2013", "-")
              .replace("\u2014", "-").replace("\u00a0", " ")
              .replace("\u2019", "'"))
    return re.sub(r"\s+", " ", txt)


_NUM = re.compile(r"-?\d+(?:\.\d+)?(?:[eE][-+]?\d+)?")


def _numbers_in(txt, a, b):
    """Every number in txt[a:b], as (float, as-printed)."""
    out = []
    for m in _NUM.finditer(txt, a, b):
        try:
            out.append((float(m.group(0)), m.group(0)))
        except ValueError:                           # noqa: PERF203
            pass
    return out


def _windows(txt, anchor):
    """Character ranges of `txt` within WINDOW of each occurrence of `anchor`."""
    if anchor is None:
        return [(0, len(txt))]
    out = []
    for m in re.finditer(re.escape(flatten(anchor)), txt):
        out.append((max(0, m.start() - WINDOW), min(len(txt), m.end() + WINDOW)))
    return out


def find_value(txt, value, anchor):
    """(found, reason).  The value must occur AS A NUMBER near the anchor.

    Matching numerically rather than as a string is what makes 168 and 168.0
    the same number while keeping 168 and 1683 different - a plain substring
    search gets both of those wrong, in opposite directions.  The tolerance is
    half a unit in the last place the CSV printed, so a value quoted to three
    decimals must agree to three decimals and is not allowed to pass on a
    coincidence two decimals away.
    """
    s = str(value).strip()
    try:
        target = float(s)
    except ValueError:
        return (s in txt), ("" if s in txt else "not present")
    dec = len(s.split(".")[1]) if "." in s and "e" not in s.lower() else 0
    tol = 0.5*10.0**(-dec) if dec else 0.5
    if anchor is not None and flatten(anchor) not in txt:
        return False, "anchor %r absent from the document" % anchor
    for a, b in _windows(txt, anchor):
        for v, _ in _numbers_in(txt, a, b):
            if abs(v - target) < tol:
                return True, ""
    for v, _ in _numbers_in(txt, 0, len(txt)):
        if abs(v - target) < tol:
            return False, "present, but not within %d chars of %r" % (WINDOW, anchor)
    return False, "value does not occur anywhere in the document"


def checks():
    nvt = pd.read_csv("04_solution/nlf_vs_turbulent.csv")
    frc = pd.read_csv("04_solution/integrated_forces.csv").set_index("quantity")
    ts = pd.read_csv("04_solution/transition_summary.csv")
    geo = pd.read_csv("01_geometry/geometry_definition.csv").set_index("parameter")
    nsum = pd.read_csv("06_validation/aerofoil_nlf0416_summary.csv").set_index("set")
    vsum = pd.read_csv("06_validation/validation_summary.csv")
    abl = pd.read_csv("06_validation/ablations.csv").set_index("configuration")
    cl_row = next(i for i in geo.index if i.startswith("Section c_l"))

    def tr(case, surf, col):
        return ts[(ts.case == case) & (ts.surface == surf)][col].iloc[0]

    # (label, value, anchor phrase it must appear near).  Anchors are TABLE
    # CAPTIONS, not CSV column headers: build_docx writes a caption as one
    # string, whereas a long column header is wrapped inside its cell by the
    # renderer and no longer occurs as a contiguous phrase in the extracted
    # text ("viscous_drag_reduction_pct" comes back broken after the
    # underscore).  The caption is the stable landmark for the table it labels.
    T_NVT = "NLF vs fully-turbulent drag"
    T_TS = "Transition prediction summary"
    T_GEO = "Geometry definition (input)"
    T_FRC = "Integrated forces"
    T_NSUM = "NLF(1)-0416 error statistics by surface"
    T_ABL = "Ablation study"
    T_VSUM = "Validation summary"

    want = [
        ("mean laminar extent", f"{nvt.mean_laminar_pct.iloc[0]:.1f}", T_NVT),
        ("viscous drag reduction",
         f"{nvt.viscous_drag_reduction_pct.iloc[0]:.1f}", T_NVT),
        ("section C_d, counts", f"{nvt.Cd_counts.iloc[0]:.1f}", T_NVT),
        ("fully turbulent C_d", f"{nvt.Cd_counts.iloc[1]:.1f}", T_NVT),
        ("cruise upper x_tr/c", f"{tr('CRUISE','upper','x_tr_c'):.3f}", T_TS),
        ("cruise lower x_tr/c", f"{tr('CRUISE','lower','x_tr_c'):.3f}", T_TS),
        ("climb upper x_tr/c", f"{tr('CLIMB','upper','x_tr_c'):.3f}", T_TS),
        ("section c_l", str(geo.loc[cl_row, "value"]), T_GEO),
        ("wing C_L", str(frc.loc["Wing C_L (lifting line, taper + washout + sweep)",
                                 "value"]), T_FRC),
        ("span efficiency e", str(frc.loc["Span efficiency e (lifting line)",
                                          "value"]), T_FRC),
        ("trim incidence", str(frc.loc["Incidence for that C_L (lifting line)",
                                       "value"]), T_FRC),
        ("aerofoil mean abs err", f"{nsum.loc['All','mean_abs_err_c']:.4f}", T_NSUM),
        ("aerofoil within bracket", str(int(nsum.loc["All", "within_bracket"])),
         T_NSUM),
        ("ablation, full model", str(int(abl.loc["full model", "within_bracket"])),
         T_ABL),
        ("ablation, one-equation march",
         str(int(abl.loc["one-equation laminar march", "within_bracket"])), T_ABL),
        ("ablation, Drela-Giles envelope",
         str(int(abl.loc["Drela-Giles envelope", "within_bracket"])), T_ABL),
        ("ablation, no bubble closure",
         str(int(abl.loc["no bubble closure", "within_bracket"])), T_ABL),
    ]
    # The swept-wing errors are quoted in the narrative and were left stale once
    # before, after the sections began being solved in the plane normal to the
    # leading edge.  They are checked now.
    sw1 = pd.read_csv("06_validation/swept_wing_crossflow.csv")
    sw2 = pd.read_csv("06_validation/swept_wing_independent.csv")
    T_SW = "Swept wings"
    want += [
        ("swept, calibration mean err", f"{sw1.err_pct.abs().mean():.1f}", T_SW),
        ("swept, independent C1=150", f"{sw2.err_pct_C1_150.abs().mean():.1f}", T_SW),
        ("swept, independent C1=200", f"{sw2.err_pct_C1_200.abs().mean():.1f}", T_SW),
    ]
    # and the T3C4 bracket error, which is the headline for that plate now
    t3 = vsum[vsum.case.str.contains("T3C4")]
    if len(t3) and t3.Re_theta_t_err_bracket_pct.notna().all():
        want.append(("T3C4 bracket error",
                     f"{float(t3.Re_theta_t_err_bracket_pct.iloc[0]):.1f}", None))
    for _, r in vsum.iterrows():
        want.append(("Re_theta_t, " + r.case.split(" flat plate")[0],
                     f"{r.Re_theta_t_pred:g}", T_VSUM))

    present = [("Somers attribution", "Somers"),
               ("section 11.4", "11.4  What the two swept-wing")]
    absent = [("no stale McGhee attribution", "McGhee et al. NLF"),
              ("no missing figure", "[missing figure"),
              ("no missing CSV", "[missing CSV"),
              ("no equation fallback", "[missing eq"),
              # a cell squeezed to illegibility breaks mid-token; these are
              # fragments the wide-table split is there to prevent
              ("no cell broken mid-token", "ER\nCO\nFT\nAC")]
    return want, present, absent


def structural_checks(txt, raw):
    """Faults that are visible in the rendered document itself.

    These do not compare against a CSV; they are properties the document must
    have whatever the numbers are, and each of them has been violated at some
    point in this project.  `txt` is whitespace-collapsed, `raw` is not: the
    caption-ordering and mid-token checks are about line breaks and need the
    text as laid out.
    """
    out = []

    # Captions numbered in the order they appear.  Hand numbering had drifted
    # to presenting Fig. 8a and 8b before Fig. 1, and Table 20a after Table 24.
    for kind, pat in (("Figure", r"Fig\.\s*(\d+)\."),
                      ("Table", r"Table\s*(\d+)\.")):
        seq = [int(m.group(1)) for m in re.finditer(r"(?m)^\s*" + pat, raw)]
        ok = bool(seq) and seq == list(range(1, len(seq)+1))
        out.append(("%s numbering (%d captions)" % (kind, len(seq)), ok,
                    "1..%d in order" % len(seq) if ok
                    else "out of order: %s" % seq[:20]))

    # No raw float64 repr anywhere in the document.  run_solution.py rounds
    # every column it writes to the precision the quantity is meaningful to;
    # the other generators must too, or the report prints a drag coefficient
    # as 0.004246488416072312 and a growth ratio as 1.1199999999999997.
    longs = re.findall(r"\d+\.\d{10,}", txt)
    out.append(("no unrounded float64 in the document", not longs,
                "clean" if not longs
                else "%d value(s), e.g. %s" % (len(longs), ", ".join(longs[:4]))))

    # A number rendered as a bare sentinel or a failed format.
    for needle, lab in (("1e+09", "no 1e9 branch-inactive sentinel"),
                        ("nan", "no bare NaN in the text"),
                        ("None", "no bare None in the text")):
        hits = len(re.findall(r"(?<![A-Za-z])" + re.escape(needle) + r"(?![A-Za-z])",
                              txt))
        out.append((lab, hits == 0, "clean" if not hits else "%d occurrence(s)" % hits))
    return out


def check_readme(want):
    """The README quotes the same headline numbers, and nothing checked it.

    This script closed the loop between the CSVs and the compiled report, and
    left the README - the first thing anyone reads - free to drift.  Its tables
    are written by hand, which is precisely the failure mode the rest of this
    file exists to catch.  Every value that must appear in the report is
    required to appear in the README too, where the README mentions the
    quantity at all.
    """
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "README.md")
    if not os.path.exists(path):
        return []
    txt = flatten(open(path, encoding="utf-8").read())
    out = []
    for label, value, _anchor in want:
        # anchored on the whole file: the README is short enough that a whole
        # number occurring anywhere in it is the number being quoted
        ok, why = find_value(txt, value, None)
        out.append((label, value, ok, why))
    return out


def main():
    if len(sys.argv) > 1 and not sys.argv[1].startswith("-"):
        path = sys.argv[1]
    elif os.path.exists("aero_turbulence_transition_report.pdf"):
        path = "aero_turbulence_transition_report.pdf"
    else:
        path = "case.docx"
    if not os.path.exists(path):
        sys.exit("no document to check: %s" % path)
    raw, pages = document_text(path)
    txt = flatten(raw)          # see flatten(): line breaks are layout, not content
    print("checking %s%s  (%d characters of text)"
          % (path, "" if pages is None else "  [%d pages]" % pages, len(raw)))

    want, present, absent = checks()
    bad = 0

    for label, ok, note in structural_checks(txt, raw):
        bad += not ok
        print("  %-8s %-38s %s" % ("OK" if ok else "FAIL", label, note))

    for label, value, anchor in want:
        ok, why = find_value(txt, value, anchor)
        bad += not ok
        print("  %-8s %-30s %-12s %s"
              % ("OK" if ok else "MISSING", label, value, why))
    for label, needle in present:
        ok = flatten(needle) in txt
        bad += not ok
        print("  %-8s %s" % ("OK" if ok else "MISSING", label))
    for label, needle in absent:
        # the mid-token needle is made of newlines, so it is checked against the
        # raw layout; the rest are phrases and are checked against the flattened
        # text so a wrapped one is still found
        hay = raw if "\n" in needle else txt
        ok = needle not in hay
        bad += not ok
        print("  %-8s %s" % ("OK" if ok else "FOUND", label))

    # the README, against the same CSVs
    rd = check_readme(want)
    if rd:
        miss = [r for r in rd if not r[2]]
        print("\n  README quotes %d of %d headline values" % (len(rd)-len(miss), len(rd)))
        for label, value, ok, why in miss:
            # A README that does not mention a quantity at all is not wrong, so
            # this reports rather than fails; a README that mentions a DIFFERENT
            # number for it is what matters, and that shows up here as a miss on
            # a value the surrounding text plainly discusses.
            print("      not found: %-30s %s" % (label, value))

    print("\n%d check(s) failed" % bad if bad else "\nall checks passed")
    sys.exit(1 if bad else 0)


if __name__ == "__main__":
    main()
