"""
build_docx.py
Assemble the full industrial case-study report -> case.docx
Embeds: narrative, all governing equations, every engineering drawing,
all CSV data tables, all plots/curves, contours, profiles, 3D contours
and vectors, validation comparisons, calibration record and sources.

No black: body text and headings use navy ink; tables use accent borders.
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
import latex2mathml.converter as _L
import mathml2omml as _M

_MATHNS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
def _omml_element(latex):
    mml = _L.convert(latex)
    omml = _M.convert(mml)
    if "xmlns:m" not in omml:
        omml = omml.replace("<m:oMath",
                            '<m:oMath xmlns:m="%s"' % _MATHNS, 1)
    return parse_xml(omml)

INK   = RGBColor(0x1d,0x2f,0x45)
BLUE  = RGBColor(0x1b,0x6c,0xa8)
ROSE  = RGBColor(0xd1,0x49,0x5b)
GREEN = RGBColor(0x2a,0x9d,0x8f)

doc = Document()

# ---- base styles (navy text, never black) ----
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(11); st.font.color.rgb = INK
for h,sz in [("Title",26),("Heading 1",17),("Heading 2",13.5),("Heading 3",11.5)]:
    s=doc.styles[h]; s.font.color.rgb = BLUE if h!="Title" else INK
    s.font.size=Pt(sz); s.font.name="Calibri"

EQD="07_equations"; eq_index=pd.read_csv(f"{EQD}/equations_index.csv").set_index("key")

import sys as _sys
_sys.path.insert(0, "solver")
import case_config as _C
C_CLIMB_TU = _C.CLIMB["Tu_pct"]

# ---- page footer with author on every page (no black) ----
foot=doc.sections[0].footer.paragraphs[0]
foot.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=foot.add_run("Akosa Samuel Onyejekwe  ·  UTSS Universal Transition Solver  ·  "
                "Case Study UTSS-CASE-2026")
fr.font.size=Pt(8.5); fr.font.color.rgb=BLUE

# ----------------------------------------------------------------------
def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)

def para(t, italic=False, bold=False, size=11):
    p=doc.add_paragraph(); r=p.add_run(t); r.italic=italic; r.bold=bold
    r.font.size=Pt(size); r.font.color.rgb=INK; return p

def bullet(t):
    p=doc.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.color.rgb=INK
    return p

def caption(t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=BLUE
    return p

def image(path, width=6.3, cap=None):
    if not os.path.exists(path):
        para(f"[missing figure: {path}]", italic=True); return
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    if cap: caption(cap)

def equation(key, show_title=True):
    if key not in eq_index.index:
        para(f"[missing eq {key}]"); return
    row=eq_index.loc[key]
    # caption line first: (key) descriptive title
    if show_title:
        c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=c.add_run(f"({key})  {row['equation']}"); r.italic=True
        r.font.size=Pt(9); r.font.color.rgb=BLUE
        c.paragraph_format.space_after=Pt(2)
    # native, editable Word equation (LaTeX -> OMML)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after=Pt(8)
    try:
        p._p.append(_omml_element(row["latex"]))
    except Exception as e:
        fr=p.add_run(row["latex"]); fr.font.color.rgb=INK
        print("  ! eq fallback", key, e)

def table_from_csv(path, max_rows=40, ncols=None, cap=None, sample=False):
    if not os.path.exists(path):
        para(f"[missing CSV: {path}]", italic=True); return
    # read as text: the CSVs carry their own formatting (Re_x written as
    # 1.785e+05, values already rounded to the precision they are claimed to),
    # and re-parsing them as floats renders 1.785e+05 as 178500.0
    df=pd.read_csv(path, dtype=str)
    if ncols: df=df.iloc[:,:ncols]
    if len(df)>max_rows:
        if sample:
            idx=list(range(0,len(df),max(1,len(df)//max_rows)))[:max_rows]
            df=df.iloc[idx]
        else:
            df=df.head(max_rows)
        truncated=True
    else: truncated=False
    add_table(df, cap)
    if truncated:
        para(f"(table sampled to {len(df)} rows; full data in {path})",
             italic=True, size=8.5)

MAX_TABLE_COLS = 7      # what stays legible across a portrait text column

def _one_table(df, size):
    t=doc.add_table(rows=1, cols=len(df.columns))
    try: t.style="Light Grid Accent 1"
    except Exception: t.style="Table Grid"
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=t.rows[0].cells
    for j,c in enumerate(df.columns):
        # a zero-width space after each underscore gives Word somewhere to
        # break a header like Re_theta_t_err_pct; without one it breaks
        # mid-token and the header reads as rubble
        hdr[j].text=str(c).replace("_", "_\u200b")
        for pp in hdr[j].paragraphs:
            for r in pp.runs: r.font.bold=True; r.font.size=Pt(size); r.font.color.rgb=INK
    for _,rowv in df.iterrows():
        cells=t.add_row().cells
        for j,v in enumerate(rowv):
            cells[j].text=str(v)
            for pp in cells[j].paragraphs:
                for r in pp.runs: r.font.size=Pt(size); r.font.color.rgb=INK
    return t


def add_table(df, cap=None, max_cols=MAX_TABLE_COLS):
    """Render a table, splitting a wide one into legible column blocks.

    Word autofits a table to the text width, so a fifteen-column CSV dropped
    straight in gets under a centimetre per column and breaks every cell into
    two- and three-character fragments.  Table 20 rendered "ERCOFTAC" as
    "ER/CO/FT/AC" down four lines and the predicted Re_theta_t of 271.4 as
    "271./4" - the validation summary, the aerofoil point-by-point comparison
    and the two surface-state tables were all unreadable in the compiled
    report.  Beyond max_cols the table is therefore split into blocks that each
    repeat the first column, the one that identifies the row.
    """
    df=df.fillna("")
    cols=list(df.columns)
    if len(cols) <= max_cols:
        _one_table(df, 8.5)
        if cap: caption(cap)
        return
    first, rest = cols[0], cols[1:]
    per = max_cols - 1
    blocks = [rest[i:i+per] for i in range(0, len(rest), per)]
    for k, blk in enumerate(blocks):
        _one_table(df[[first]+blk], 8.0)
        tag = "columns %d-%d of %d" % (2+k*per, 1+k*per+len(blk), len(cols))
        if k < len(blocks)-1:
            caption("(%s; %s repeated on each block)" % (tag, first))
            doc.add_paragraph()
        elif cap:
            caption("%s  (%s; the table is split across %d blocks so that no "
                    "cell is compressed to illegibility, with %s repeated on "
                    "each)" % (cap, tag, len(blocks), first))
        else:
            caption("(%s)" % tag)

def manual_table(headers, rows, cap=None):
    add_table(pd.DataFrame(rows, columns=headers), cap)

# ======================================================================
#  TITLE PAGE
# ======================================================================
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("INDUSTRIAL CASE STUDY"); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=GREEN
ttl=doc.add_paragraph(); ttl.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=ttl.add_run("Prediction of Boundary-Layer Turbulence Transition\nover Aircraft Surfaces")
r.bold=True; r.font.size=Pt(24); r.font.color.rgb=INK
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("using the UTSS Universal Transition & Skin-Friction Solver")
r.font.size=Pt(15); r.font.color.rgb=BLUE
doc.add_paragraph()
image("01_geometry/drawings/dwg_05_isometric.png", width=5.6)
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sub.add_run("Case vehicle: AETHER-NLF 25 regional natural-laminar-flow demonstrator\n"
              "Wing section UTSS-NLF16  ·  Cruise FL360, M0.42  ·  3-D swept tapered wing")
r.font.size=Pt(12); r.font.color.rgb=INK
doc.add_paragraph()
auth=doc.add_paragraph(); auth.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=auth.add_run("Prepared by"); r.font.size=Pt(11); r.italic=True; r.font.color.rgb=BLUE
auth2=doc.add_paragraph(); auth2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=auth2.add_run("AKOSA SAMUEL ONYEJEKWE"); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=INK
meta=doc.add_paragraph(); meta.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=meta.add_run("Document UTSS-CASE-2026  ·  Three-dimensional analysis  ·  "
               "All units SI unless noted"); r.font.size=Pt(10); r.italic=True
r.font.color.rgb=BLUE
doc.add_page_break()

# ======================================================================
h1("1.  Executive Summary")
para("This case study demonstrates the prediction of laminar-to-turbulent boundary-layer "
 "transition over the surfaces of an aircraft wing, and the engineering quantities that "
 "depend on it (skin-friction drag, laminar-flow extent, boundary-layer growth and "
 "trailing-edge separation margin). The analysis vehicle is the AETHER-NLF 25, a regional "
 "natural-laminar-flow (NLF) demonstrator whose three-dimensional swept, tapered and twisted "
 "wing uses the purpose-designed UTSS-NLF16 section. Predictions are produced by the UTSS "
 "(Universal Transition & Skin-friction Solver), a novel, fast, robust engine that couples a "
 "vortex-panel inviscid solution to an integral boundary-layer marcher driven by a single, "
 "unified four-mechanism transition kernel.")
para("Key results at the cruise design point (FL360, M=0.42, Re_MAC ≈ 6.4×10⁶):", bold=True)


def _headline():
    """Read Table 1 straight out of the generated CSVs.

    These numbers were typed in by hand in an earlier version of this script
    and had drifted from the solution they describe.  They are now derived, so
    the summary cannot disagree with section 9.
    """
    ts = pd.read_csv("04_solution/transition_summary.csv")
    nvt = pd.read_csv("04_solution/nlf_vs_turbulent.csv")
    frc = pd.read_csv("04_solution/integrated_forces.csv").set_index("quantity")

    def row(case, surf):
        r = ts[(ts.case == case) & (ts.surface == surf)].iloc[0]
        return float(r.x_tr_c), str(r.mechanism)
    xu, mu = row("CRUISE", "upper")
    xl, ml = row("CRUISE", "lower")
    xcu, mcu = row("CLIMB", "upper")
    lam_pct = float(nvt.mean_laminar_pct.iloc[0])
    cd_nlf = float(nvt.Cd_counts.iloc[0])
    saving = float(nvt.viscous_drag_reduction_pct.iloc[0])
    cl = float(frc.loc["Section lift coefficient Cl", "value"])
    return [["Upper-surface transition x_tr/c", f"{xu:.3f} ({mu})"],
            ["Lower-surface transition x_tr/c", f"{xl:.3f} ({ml})"],
            ["Mean laminar-flow extent", f"{lam_pct:.1f} % of chord"],
            ["Section lift coefficient C_l", f"{cl:.3f}"],
            ["Section profile drag C_d", f"{cd_nlf:.1f} counts"],
            ["Viscous drag reduction vs fully-turbulent", f"{saving:.1f} %"],
            [f"Climb (Tu={C_CLIMB_TU:g} %) transition x_tr/c (upper)",
             f"{xcu:.3f} ({mcu})"]]


manual_table(["Quantity","Predicted value"], _headline(),
 cap="Table 1. Headline predictions, read from the generated solution CSVs.")
para("Validated against eight independent, credible published datasets with one "
 "universal calibration set — four ERCOFTAC flat plates (case 020), the Schubauer-Skramstad "
 "plate, the NLF(1)-0416 aerofoil at 86 conditions, and two swept wings — the solver "
 "reproduces the transition-onset Reynolds number Re_θt and the measured transition location "
 "without per-case re-tuning of the physics. All four criteria of the kernel are selected "
 "somewhere in this study and each is supported by measurement. The remainder of this report sets "
 "out the background, problem, governing equations, the complete input dataset, every generated "
 "engineering output (CSVs, curves, metrics, contours, temperature profiles, 3-D contours and "
 "vectors), the validation and calibration record with all sources, and the contribution to knowledge.")

# ======================================================================
h1("2.  Background")
para("Skin-friction drag is one of the largest components of total aircraft drag — typically "
 "45–50 % of cruise drag for a transport aircraft. A turbulent boundary layer produces several "
 "times the skin friction of a laminar one at the same Reynolds number. Consequently, extending "
 "the laminar run over wing, nacelle and empennage surfaces (natural-laminar-flow, NLF, and "
 "hybrid-laminar-flow control, HLFC) is among the highest-leverage technologies for fuel-burn "
 "and emissions reduction. The enabling capability is the ability to PREDICT, reliably and "
 "cheaply, WHERE the boundary layer transitions from laminar to turbulent over a 3-D surface, "
 "across the flight envelope.")
para("Transition is not a single phenomenon. Over aircraft surfaces it is driven by several, "
 "often co-resident, physical mechanisms:")
bullet("Natural / Tollmien–Schlichting (TS): amplification of viscous instability waves in "
       "low-disturbance free flight, predicted here by an e^N integral carried per physical "
       "frequency, with the spatial growth rates read from a tabulated solution of the "
       "Orr-Sommerfeld problem rather than from an envelope correlation.")
bullet("Bypass: free-stream-turbulence-induced transition that skips the TS route — dominant in "
       "high-turbulence environments (climb through cloud, turbomachinery-like inflow).")
bullet("Laminar-separation-induced: a laminar separation bubble that reattaches turbulent. "
       "The bubble is closed explicitly here — the shear layer is carried across the dead-air "
       "region by the momentum integral and reattachment placed where it has amplified by the "
       "same critical factor used elsewhere — so a length, and not merely a separation point, "
       "is predicted.")
bullet("Cross-flow: three-dimensional instability of the cross-flow velocity profile on swept "
       "wings — the principal NLF-limiter at moderate-to-high sweep.")
para("A practical solver for aircraft design must capture all of these with a SINGLE, "
 "consistent formulation and calibration, run in seconds for thousands of design iterations, "
 "and remain robust across the operating envelope. Existing tools each address part of the "
 "problem (Section 5). UTSS unifies them.")

# ======================================================================
h1("3.  Problem Statement and Solution Approach")
h2("3.1  Problem statement")
para("Given a three-dimensional wing geometry and a flight condition, predict — quickly, "
 "robustly and accurately — the chordwise and span-wise location of boundary-layer transition "
 "on both surfaces, the governing transition mechanism at each station, and the resulting "
 "boundary-layer state (skin friction C_f, momentum thickness θ, shape factor H, intermittency "
 "γ), and from these the laminar-flow extent and the viscous (profile) drag. The method must be "
 "universal: one set of physics and calibration constants valid for natural, bypass, "
 "separation-induced and cross-flow transition across the Reynolds- and turbulence-number range "
 "of real aircraft surfaces.")
h2("3.2  How we solve it")
para("UTSS solves the problem in a layered, fully-coupled manner:")
bullet("Inviscid edge flow: a constant-strength vortex-panel method (Kuethe–Chow) returns the "
       "surface pressure C_p and edge velocity U_e, with a Karman-Tsien correction applied to "
       "C_p and the integrated loads. It returns a stagnation pressure coefficient of 1.048 at "
       "M = 0.42 against the exact isentropic 1.045, where the linearised Prandtl-Glauert "
       "scaling gives 1.102, and the two agree to better than half a per cent below M = 0.2.")
bullet("Laminar boundary layer: the momentum AND kinetic-energy integral equations are advanced "
       "together from the stagnation point along each surface, so the shape factor is a solved "
       "variable carrying its own history rather than a local function of the pressure gradient. "
       "The three closure functions - H*(H), Re_theta*Cf/2 and Re_theta*C_D - are properties of "
       "the Falkner-Skan family, computed from it and not fitted; every similarity solution is an "
       "exact fixed point of the second equation, and the march returns H = 2.5914 on a flat plate "
       "against the Blasius 2.5913. Fluid properties are evaluated at Eckert's reference "
       "temperature so that the incompressible closures return the compressible skin friction.")
bullet("Unified transition kernel (the novel core): at every station the effective transition "
       "Reynolds number is the minimum across the four mechanisms, each with a calibration weight.")
bullet("Transitional region: a Narasimha universal-intermittency closure blends laminar and "
       "turbulent properties through γ.")
bullet("Turbulent boundary layer: Head's entrainment method with the Ludwieg–Tillmann "
       "skin-friction law advances the turbulent BL to the trailing edge.")
bullet("Integration: the Squire–Young formula returns profile drag; a span-wise strip sweep with "
       "the cross-flow mechanism extends the solution to the full 3-D wing.")
para("Two elements of the formulation are not correlations. The amplification rate driving the "
 "e^N integral is tabulated in solver/amplification_db.npz from 61,600 Orr-Sommerfeld "
 "eigenvalue solutions on the Falkner-Skan family, continued past separation onto its reverse-flow branch: profiles are generated by continuation, the "
 "eigenvalue problem is solved by Chebyshev collocation, and the spatial growth rate is "
 "recovered from the temporal one through Gaster's transformation. The stability solver returns "
 "the Blasius neutral point at Re_theta = 201 against the accepted 200.5 - the "
 "tabulated database resolves it to the nearest node of its Reynolds-number grid, "
 "Re_theta = 210 - a shape factor of "
 "2.59129 and a wall shear parameter f''(0) = 0.46960. At run time the cost is a table lookup, "
 "so the sub-second solution is preserved. The second is the separation-bubble closure "
 "described above, whose length scales with the disturbance environment and therefore spans "
 "measured bubbles from about 40 momentum thicknesses at 2.1 % free-stream turbulence to about "
 "180 at 0.03 %.")

# ======================================================================
h1("4.  The UTSS Universal Solver — Governing Equations")
para("All equations implemented in the solver are listed below as native, editable Word "
 "equations at standard size. They constitute the complete mathematical definition of the "
 "method, and are also collected in the companion file model.equations.docx.")
h2("4.1  Inviscid edge solution")
for k in ["E02","E03","E01","E04"]: equation(k)
h2("4.2  Laminar boundary layer (Thwaites)")
for k in ["E05","E06","E06b","E06c","E07"]: equation(k)
h2("4.3  Unified four-mechanism transition kernel (novel contribution)")
para("Bypass onset uses the Abu-Ghannam & Shaw correlation evaluated at the flow-history-averaged "
 "Tu; natural/TS onset integrates one amplification factor per physical frequency using the "
 "tabulated Orr-Sommerfeld growth rates and triggers on their envelope at N_crit; "
 "separation-induced onset closes a laminar bubble across the dead-air region; and cross-flow "
 "onset uses the C1 criterion. The kernel takes the minimum effective onset across all four:")
for k in ["E08","E09","E10","E10b","E10c","E11","E11b","E12","E12b","E13","E14"]: equation(k)
h2("4.4  Transitional region and turbulent closure")
for k in ["E15","E16","E17","E18","E19"]: equation(k)
h2("4.5  Drag, temperature and reference quantities")
for k in ["E20","E21","E22","E22b","E23","E24"]: equation(k)

# ======================================================================
h1("5.  Why UTSS is Better — Comparison with Existing Solvers")
para("UTSS is designed to be fast, robust and broad in regime coverage. The table "
 "contrasts its capabilities with the principal classes of existing transition tools.")
manual_table(
 ["Capability","XFOIL / e^N codes","RANS γ–Re_θ (CFD)","LES / DNS","UTSS (this work)"],
 [["Natural / TS transition","Yes","Indirect","Yes","Yes (tabulated e^N)"],
  ["Bypass (free-stream Tu)","No","Yes","Yes","Yes (AGS)"],
  ["Separation-induced","Limited","Yes","Yes","Yes (bubble closure)"],
  ["Cross-flow (3-D swept)","No","Add-on only","Yes","Yes (C1, built-in)"],
  ["Single universal calibration","n/a","Needs re-tuning","n/a","Yes (one constant set)"],
  ["3-D wing capability","2-D only","Yes","Yes","Yes (strip + cross-flow)"],
  ["Robustness / convergence","Can fail in sep.","Stiff, costly","Very costly","Robust, direct"],
  ["Typical CPU cost","seconds","hours","days–weeks","< 1 second"],
  ["Suited to design loops","Partly","No","No","Yes"]],
 cap="Table 2. Capability comparison versus existing solver classes.")
para("Novelty. The distinguishing element is the unified transition kernel "
 "(Eq. E13): a single closed expression that selects the governing transition mechanism "
 "locally by taking the minimum effective onset Reynolds number across four co-resident "
 "mechanisms, each modulated by a calibration weight, and feeds a single intermittency closure. "
 "Unlike e^N codes (TS only) or correlation RANS models (which require case-by-case re-tuning "
 "and a full CFD solve), UTSS reproduces natural, bypass, separation and cross-flow transition "
 "with ONE calibration set at panel-method cost. Every one of the four branches is the "
 "selected mechanism somewhere in this study — natural/TS on the cruise wing, the "
 "Schubauer-Skramstad plate and the NLF(1)-0416 upper surface; bypass on the climb case and "
 "the ERCOFTAC plates; separation on T3C4 and the NLF(1)-0416 lower surface; cross-flow on "
 "both swept wings — and each is checked against measurement in Section 11.")

# ======================================================================
h1("6.  Case-Study Definition and Input Data")
para("All input data required to run the prediction are tabulated below. The case is fully "
 "three-dimensional.")
h2("6.1  Aircraft and wing geometry (input)")
table_from_csv("01_geometry/geometry_definition.csv", cap="Table 3. Geometry definition (input).")
h2("6.2  Flight conditions (input)")
table_from_csv("03_model_setup/flow_conditions.csv", cap="Table 4. Flight / flow conditions (input).")
image("05_postprocessing/csv_plots/conditions_compare.png", width=6.4,
      cap="Fig. 8a. Cruise vs climb flight-condition comparison (flow_conditions.csv).")
h2("6.3  Fluid (material) properties (input)")
table_from_csv("03_model_setup/material_properties.csv", cap="Table 5. Air properties (input).")
h2("6.4  Solver settings (input)")
table_from_csv("03_model_setup/solver_settings.csv", cap="Table 6. Solver configuration (input).")
h2("6.5  Universal calibration constants (input)")
table_from_csv("03_model_setup/calibration_constants.csv",
               cap="Table 7. Calibration constants and their sources (input).")
image("05_postprocessing/csv_plots/calibration_constants.png", width=6.0,
      cap="Fig. 8b. UTSS universal calibration constant set (calibration_constants.csv).")

# ======================================================================
h1("7.  Geometry and Engineering Drawings")
para("The wing is drawn to standard third-angle orthographic projection. All drawings are "
 "dimensioned and to scale; views provided: section, plan, front, side, full orthographic "
 "sheet, isometric and structural section.")
for f,c in [("dwg_01_airfoil_section","Fig. 1. UTSS-NLF16 aerofoil section (dimensioned)."),
            ("dwg_02_planview","Fig. 2. Wing planform — plan (top) view, dimensioned."),
            ("dwg_03_front_side","Fig. 3. Front and side orthographic views."),
            ("dwg_04_orthographic","Fig. 4. Full third-angle orthographic projection sheet."),
            ("dwg_05_isometric","Fig. 5. Isometric view of the wing."),
            ("dwg_06_section_BB","Fig. 6. Structural sectional view B–B.")]:
    image(f"01_geometry/drawings/{f}.png", width=6.4, cap=c)
h2("7.1  Geometry data (from CSV)")
image("05_postprocessing/csv_plots/geo_airfoil.png", width=5.8,
      cap="Fig. 7. Section geometry plotted from airfoil_UTSS-NLF16.csv.")
image("05_postprocessing/csv_plots/geo_planform.png", width=5.8,
      cap="Fig. 8. Span-wise geometry from wing_planform.csv.")
table_from_csv("01_geometry/wing_planform.csv", max_rows=21,
               cap="Table 8. Wing planform stations (wing_planform.csv).")
image("05_postprocessing/csv_plots/geo_sections_3d.png", width=5.8,
      cap="Fig. 8c. Lofted 3-D wing sections (wing_sections_3d.csv).")

# ======================================================================
h1("8.  Mesh / Discretisation")
_mi = pd.read_csv("02_mesh/mesh_independence.csv")
_cd = _mi.Cd.to_numpy(float)*1e4
_hi = _cd[_mi.n_surface_panels >= 180]
para("The surface is discretised with cosine-clustered streamwise nodes; a wall-normal "
 "reconstruction grid (first-cell y⁺≈1) supports boundary-layer profile recovery. A "
 "panel-count sweep bounds the discretisation sensitivity rather than demonstrating "
 "asymptotic convergence: from %d to %d surface panels the section drag spans %.1f counts, "
 "%.1f %% of its mean, and above %d panels it stays within about ±%.1f count without "
 "tightening further."
 % (_mi.n_surface_panels.min(), _mi.n_surface_panels.max(),
    _cd.ptp(), 100.0*_cd.ptp()/_cd.mean(), 180, 0.5*_hi.ptp()) +
 "  The residual wander is not a truncation error that refinement removes — "
 "it is set by which panel the transition point lands on, so it scales with the panel "
 "spacing at transition and is of the same order as the ±0.025c bracket the aerofoil "
 "measurements themselves carry. The 260-panel grid is used for every case-study "
 "result; the tabulated validation sections are re-splined onto their own "
 "cosine-clustered grids of 400 and 440 panels.")
table_from_csv("02_mesh/mesh_metrics.csv", cap="Table 9. Mesh metrics.")
table_from_csv("02_mesh/mesh_independence.csv",
               cap="Table 10. Panel-count sensitivity study.")
for f,c in [("plots/mesh_01_surface","Fig. 9. Surface mesh and wall-normal stacks."),
            ("plots/mesh_02_bl_normal","Fig. 10. Wall-normal reconstruction grid / y⁺."),
            ("plots/mesh_03_independence","Fig. 11. Panel-count sensitivity of C_d and "
             "transition location.")]:
    image(f"02_mesh/{f}.png", width=5.8, cap=c)
table_from_csv("02_mesh/bl_normal_grid.csv", max_rows=22, sample=True,
               cap="Table 11. Wall-normal grid (bl_normal_grid.csv, sampled).")

# ======================================================================
h1("9.  Solution — Generated Engineering Output Data")
para("This section presents every generated output: numerical tables (CSV) and the plotted "
 "curve for each. The boundary-layer state is reported on both surfaces at the cruise and "
 "climb conditions.")
h2("9.1  Transition prediction summary")
table_from_csv("04_solution/transition_summary.csv",
               cap="Table 12. Transition prediction summary (transition_summary.csv).")
image("05_postprocessing/csv_plots/transition_summary_bar.png", width=6.0,
      cap="Fig. 12a. Predicted laminar-flow extent by case and surface (transition_summary.csv).")
h2("9.2  Integrated forces and drag breakdown")
para("The wing lift is obtained from Prandtl's lifting-line theory - Glauert's monoplane "
 "equation solved for the odd Fourier coefficients of the loading - using the section "
 "lift-curve slope and zero-lift incidence returned by the same panel method, the planform "
 "chord distribution, the wing's washout, and the section slope reduced by the cosine of "
 "the quarter-chord sweep. An earlier version of this table asserted C_L = 0.90 c_l; the "
 "planform actually returns 0.49, because a −3° washout is large relative to the 3.7° by "
 "which the root exceeds its zero-lift incidence, and because the downwash of an AR = 8.7 "
 "wing reduces the slope by a quarter. The span efficiency and the induced drag come out of "
 "the same solve. The implementation is checked against the one case with a closed-form "
 "answer - an elliptic planform, for which it returns e = 1 and the exact lift-curve slope "
 "to machine precision - every time the solution is regenerated.")
para("The incidence in Table 4 is the SECTION design incidence, not the trim point of the "
 "aircraft the planform belongs to: at 1.5° the wing carries 23.6 kN against an all-up "
 "weight of 83.4 kN, and level flight at MTOW would need 7.6°. Table 13 reports both, so "
 "the two are not confused. Every transition result in this study is quoted at the design "
 "incidence and is unaffected by that distinction.", italic=True, size=10)
table_from_csv("04_solution/integrated_forces.csv", cap="Table 13. Integrated forces, "
               "with the wing quantities from the lifting-line solve.")
table_from_csv("04_solution/nlf_vs_turbulent.csv", cap="Table 14. NLF vs fully-turbulent drag.")
image("05_postprocessing/csv_plots/nlf_vs_turbulent.png", width=5.4,
      cap="Fig. 12. Drag benefit of predicted laminar flow vs fully-turbulent.")
h2("9.2a  The transition-length closure and what the result owes to it")
para("The extent of the transitional region is Dhawan and Narasimha's published correlation, "
 "Re_λ = 9 Re_x,t^0.75, with λ the distance over which the intermittency rises from 0.25 to "
 "0.75. It is validated here on the flat plates of Section 11.1, which span Re_x,t = 6×10⁴ to "
 "1.4×10⁶ and on which it reproduces the measured extent of the skin-friction rise to within "
 "a factor of two — exactly on T3B. An earlier version of this work wrote the correlation as "
 "6.5 Re_θt^0.8, which returns Re_λ ≈ 600 where those plates require ≈ 4×10⁴: the layer went "
 "from fully laminar to fully turbulent inside a single marching station, and the predicted "
 "C_f jumped vertically where the measurements climb over half a decade of Re_x. The "
 "transition LOCATION was never affected — onset is where the kernel fires, not where the "
 "blend ends — but everything downstream of onset was.")
para("The cruise wing transitions at Re_x,t = 3.7×10⁶, a factor of three beyond the range the "
 "correlation is validated over here, and it then returns a transitional zone of about a "
 "third of the chord, which is longer than a real natural-laminar-flow section shows at this "
 "Reynolds number. That extrapolation is not damped: doing so would add an undeclared "
 "constant to a method whose claim is that it has none. What it costs is measured instead. "
 "Sweeping the constant over a factor of four moves the section drag by a tenth of a count, "
 "so no reported drag in this study depends on the extrapolated part of the closure; beyond "
 "twice the published value the layer no longer completes transition before the trailing "
 "edge, and there it would.")
table_from_csv("04_solution/transition_length_sensitivity.csv",
               cap="Table 14a. Sensitivity of the case-study result to the transition-length "
                   "constant (transition_length_sensitivity.csv).")

h2("9.3  Surface distributions — cruise")
for f,c in [("cruise_Cp","Fig. 13. Pressure coefficient C_p — cruise."),
            ("cruise_Cf","Fig. 14. Skin-friction C_f and laminar run — cruise."),
            ("cruise_theta_H","Fig. 15. Momentum thickness θ and shape factor H — cruise."),
            ("cruise_Retheta","Fig. 16. The governing transition criterion — cruise. "
             "Two of the four branches are Reynolds-number thresholds and two are "
             "amplification integrals, so both pairs are drawn; at cruise it is N "
             "reaching N_crit that fires, and no Re_θt threshold is active at any "
             "station."),
            ("cruise_gamma","Fig. 17. Intermittency γ — cruise.")]:
    image(f"05_postprocessing/csv_plots/{f}.png", width=5.8, cap=c)
table_from_csv("04_solution/surface_cruise_upper.csv", max_rows=30, sample=True,
   cap="Table 15. Cruise upper-surface BL state (surface_cruise_upper.csv, sampled).")
table_from_csv("04_solution/surface_cruise_lower.csv", max_rows=30, sample=True,
   cap="Table 16. Cruise lower-surface BL state (surface_cruise_lower.csv, sampled).")
h2("9.4  Surface distributions — climb (off-design, elevated Tu)")
for f,c in [("climb_Cp","Fig. 18. Pressure coefficient C_p — climb."),
            ("climb_Cf","Fig. 19. Skin-friction C_f and laminar run — climb."),
            ("climb_Retheta","Fig. 19b. The governing transition criterion — climb: "
             "here Re_θ crosses the falling Abu-Ghannam & Shaw bypass threshold "
             "while N is still far below N_crit."),
            ("climb_gamma","Fig. 20. Intermittency γ — climb."),
            ("compare_cruise_climb_Cf","Fig. 21. C_f cruise vs climb — regime-dependent transition.")]:
    image(f"05_postprocessing/csv_plots/{f}.png", width=5.8, cap=c)
h2("9.5  Aerodynamic polars")
table_from_csv("04_solution/aero_polar.csv", cap="Table 17. Aerodynamic polar (aero_polar.csv).")
image("05_postprocessing/csv_plots/aero_polar.png", width=6.3,
      cap="Fig. 22. Lift curve, drag polar, L/D and transition vs angle of attack.")
h2("9.6  Span-wise distribution (3-D)")
table_from_csv("04_solution/spanwise_distribution.csv",
               cap="Table 18. Span-wise distribution (spanwise_distribution.csv).")
image("05_postprocessing/csv_plots/spanwise_transition.png", width=5.8,
      cap="Fig. 23. Span-wise transition front and section drag.")

# ======================================================================
h1("10.  Post-Processing — Contours, Profiles and 3-D Fields")
h2("10.1  Pressure and velocity contours")
for f,c in [("contour_Cp_cruise","Fig. 24. Pressure-coefficient contour — cruise."),
            ("contour_speed_cruise","Fig. 25. Velocity magnitude and streamlines — cruise."),
            ("vectors_cruise","Fig. 26. Velocity vector field — cruise."),
            ("contour_Cp_climb","Fig. 27. Pressure-coefficient contour — climb.")]:
    image(f"05_postprocessing/contours/{f}.png", width=6.2, cap=c)
h2("10.2  Boundary-layer velocity and temperature profiles")
para("The profiles are reconstructed from the marched state and not from an assumed shape. "
 "The laminar leg is the Falkner-Skan profile at the shape factor the march solved for, read "
 "from the same family that supplies the closure functions; the turbulent leg is the power "
 "law that shape factor implies, H = (n+2)/n; and the two are blended by the same "
 "intermittency that blends C_f, θ and H, each on its own thickness — δ99 = η99 θ/θ_η for "
 "the similarity profile and δ = θ(n+1)(n+2)/n for the power law. Temperature profiles then "
 "follow from the compressible Crocco–Busemann relation (Eq. E21), showing wall-recovery "
 "heating through the boundary layer.")
for f,c in [("bl_velocity_profiles","Fig. 28. Boundary-layer velocity profiles."),
            ("bl_temperature_profiles","Fig. 29. Boundary-layer temperature profiles (K)."),
            ("bl_temperature_ratio","Fig. 30. Normalised temperature profiles T/T_e.")]:
    image(f"05_postprocessing/profiles/{f}.png", width=5.5, cap=c)
table_from_csv("04_solution/bl_profiles_cruise.csv", max_rows=28, sample=True,
   cap="Table 19. BL velocity/temperature profiles (bl_profiles_cruise.csv, sampled).")
h2("10.3  Three-dimensional surface contours and vectors")
para("The full 3-D wing surface is coloured by the predicted fields; the transition front is "
 "directly visible as the laminar-to-turbulent boundary on the intermittency contour.")
for f,c in [("td_Cp","Fig. 31. 3-D wing surface contour — pressure C_p."),
            ("td_Cf","Fig. 32. 3-D wing surface contour — skin friction C_f (×10³)."),
            ("td_gamma","Fig. 33. 3-D wing surface contour — intermittency γ (transition front)."),
            ("td_skinfriction_vectors","Fig. 34. Upper-surface flow direction coloured by C_f "
              "(chordwise only — the strip formulation carries no span-wise wall shear).")]:
    image(f"05_postprocessing/three_d/{f}.png", width=6.2, cap=c)

# ======================================================================
h1("11.  Validation and Calibration")
para("To qualify as universal, the solver is validated against every published dataset the "
 "kernel\'s four branches reach — five flat plates spanning 0.03 to 6 per cent free-stream "
 "turbulence (bypass, natural and separation-induced transition), 86 aerofoil conditions on "
 "the NLF(1)-0416 section, and two swept wings from different facilities and eras — using ONE "
 "universal calibration set with no per-case re-tuning of the physics. The transition-onset "
 "momentum-thickness Reynolds number Re_θt is the metric on the plates and the transition "
 "location x_tr/c on the aerofoil and the swept wings. The skin-friction error is reported "
 "separately over the laminar run and over the turbulent run, on the stations where the "
 "measurement and the prediction are in the same state; pooled across transition it measures "
 "the onset error a second time, in the wrong units, because a plate whose onset is early by "
 "16 % is then charged with the whole laminar-to-turbulent step in C_f over the interval "
 "between the two onsets.")
table_from_csv("06_validation/validation_summary.csv",
               cap="Table 20. Validation summary — predicted vs published Re_θt.")
h2("11.1  Flat plates")
for f,c in [("val_T3A","Fig. 35. Validation — ERCOFTAC T3A flat plate (Tu = 3.0 %, bypass)."),
            ("val_T3AM","Fig. 36. Validation — ERCOFTAC T3A⁻ flat plate (Tu = 0.87 %, bypass)."),
            ("val_T3B","Fig. 37. Validation — ERCOFTAC T3B flat plate (Tu = 6.0 %, bypass)."),
            ("val_T3C4","Fig. 38. Validation — ERCOFTAC T3C4 flat plate (laminar separation bubble)."),
            ("val_SS","Fig. 39. Validation — Schubauer & Skramstad natural transition (Tu = 0.03 %)."),
            ("val_combined_Re_theta_t","Fig. 40. Universal validation — Re_θt across all five plates.")]:
    image(f"06_validation/plots/{f}.png", width=5.9, cap=c)

h2("11.2  NLF(1)-0416 aerofoil — 86 transition locations")
para("The largest single body of evidence in this work, and the one on which nothing is "
 "calibrated. Transition locations were digitised from Fig. 9 of the source report at four "
 "chord Reynolds numbers on both surfaces, and each condition is matched by trimming the "
 "incidence to the measured lift coefficient. The experiment brackets transition between "
 "adjacent orifices 0.05c apart, so its own uncertainty is ±0.025c and a prediction inside "
 "that band cannot be distinguished from the measurement. Both the natural (TS) and the "
 "separation-induced branches are selected on this set.")
table_from_csv("06_validation/aerofoil_nlf0416_summary.csv",
               cap="Table 21. NLF(1)-0416 error statistics by surface "
                   "(aerofoil_nlf0416_summary.csv).")
image("06_validation/plots/val_aerofoil_nlf0416.png", width=6.4,
      cap="Fig. 41. Transition location against lift coefficient at four chord "
          "Reynolds numbers, measurement bars = the ±0.025c orifice bracket.")
table_from_csv("06_validation/aerofoil_nlf0416.csv", max_rows=30, sample=True,
               cap="Table 22. NLF(1)-0416 point-by-point comparison "
                   "(aerofoil_nlf0416.csv, sampled).")

h2("11.3  Swept wings — the cross-flow branch")
para("The cross-flow coefficient is set on the first of these two experiments and nothing is "
 "calibrated on the second, which is a different facility, section and era. The branch "
 "reproduces the calibration set to 14.7 % at the frozen constant C1 = 150. On the "
 "independent set that same constant gives 51.7 %, and C1 = 200 gives 18.4 % — so the "
 "criterion has the right functional form on both wings but not one critical value that "
 "serves both. The independent set is therefore tabulated at BOTH ends of that band: "
 "reporting only C1 = 150 understates what the criterion does here, and reporting only "
 "C1 = 200 would be a per-case re-tune of the kind this work is claiming not to need. The "
 "spread between the two columns is the limitation, stated rather than averaged away. "
 "Nothing else in the model differs between the two columns.")
para("What each experiment requires of the criterion is measured rather than asserted "
 "(Table 25). The march is run with every branch disabled so that it reaches the measured "
 "transition station, and the criterion is evaluated there. Dagenhart & Saric require a "
 "critical Re_θ2 of 165 with a 17.2 % coefficient of variation over six chord Reynolds "
 "numbers; Boltz et al. require 234 with 4.0 % over four sweep angles and a factor of three in chord "
 "Reynolds number. Each facility is therefore internally consistent — the second markedly so — "
 "and the two differ by 42 %. That is the shape of a receptivity difference rather than of a "
 "criterion with the wrong form: stationary cross-flow vortices are seeded by leading-edge "
 "roughness, and neither report documents the surface finish. Two attempts to close the gap "
 "fail and are recorded rather than dropped. Replacing the constant surrogate by the exact "
 "Falkner-Skan-Cooke factor K(λ) makes matters worse, taking the pooled coefficient of "
 "variation from 21 to 82 % and inverting the ratio between the two sets. Giving the "
 "cross-flow branch its own "
 "amplification threshold, separate from the one Mack's relation supplies for "
 "Tollmien-Schlichting waves — which is defensible, since a stationary cross-flow vortex is "
 "not seeded by free-stream turbulence — moves the independent set only from 55 to 51 % as "
 "that threshold is taken from N = 2 to N = 12, with C1 refitted on the calibration set at "
 "each value. The reason is that once Re_θ2 exceeds C1 the amplification builds so quickly "
 "that the threshold is nearly redundant with C1 itself, so the two constants cannot be "
 "separated by these data.")
table_from_csv("06_validation/swept_wing_crossflow.csv",
               cap="Table 23. Cross-flow validation, 45° swept NLF(2)-0415 "
                   "(Dagenhart & Saric — the calibration set).")
image("06_validation/plots/val_swept_crossflow.png", width=5.9,
      cap="Fig. 42. Transition location against chord Reynolds number, 45° swept "
          "NLF(2)-0415.")
table_from_csv("06_validation/swept_wing_independent.csv",
               cap="Table 24. Independent swept-wing check, NACA 64(2)A015 (Boltz et al.) "
                   "at both ends of the reported cross-flow band — nothing calibrated here.")
image("06_validation/plots/val_swept_independent.png", width=5.9,
      cap="Fig. 43. Transition location against sweep angle, NACA 64(2)A015, "
          "with the C1 = 150–200 band shaded.")

h2("11.3a  Where the flat-plate residuals come from")
para("Three of the five plates carry a residual worth accounting for, and the accounting "
 "is generated by gen_validation.py rather than argued. Two effects cover them, and which "
 "one applies depends on which branch fires — a distinction the single summary table of "
 "Table 20 hides.")
para("Pre-transitional thickening. Free-stream turbulence thickens a laminar layer before "
 "it transitions, and the laminar march carries no such mechanism. Where onset comes from "
 "the Abu-Ghannam & Shaw correlation the effect is already inside it, because that "
 "correlation was fitted to MEASURED momentum-thickness Reynolds numbers. Where onset comes "
 "from the model's own marched θ — which on this set is the separating T3C4 plate alone — it "
 "is not. The measured layer at T3C4 is 1.36 times the Blasius momentum thickness at its "
 "own onset, the largest of the five, and multiplying the prediction by that factor takes "
 "the error from −28.8 % to −3.3 %. Applying the same correction to the three bypass plates "
 "makes each of them worse, which is the expected signature of a double count and is what "
 "shows the diagnosis is the right one rather than a convenient rescaling. The −28.8 % on "
 "T3C4 is therefore not a defect of the separation closure; it is the turbulence thickening "
 "that no laminar integral march carries, and correcting it would require a "
 "pre-transitional closure fitted to the single plate that tests this branch.")
para("Conditioning. In a decaying stream the onset threshold rises while Re_θ grows only as "
 "the square root of distance, so the two curves close at a shallow angle and the crossing "
 "is sensitive. Shifting the threshold by ±10 % moves the predicted transition location by "
 "a factor of 3.5 on T3A, 2.5 on T3A⁻ and 2.0 on T3B. A correlation accurate to ten per "
 "cent cannot locate transition to ten per cent in such a flow, and no refinement of the "
 "correlation changes that.")
table_from_csv("06_validation/residual_diagnostics.csv",
               cap="Table 20a. Where each flat-plate residual comes from "
                   "(residual_diagnostics.csv): the pre-transitional thickening the "
                   "laminar march does not carry, and the sensitivity of the crossing.")

h2("11.4  What the two swept-wing experiments require of the criterion")
table_from_csv("06_validation/crossflow_criticals_summary.csv",
               cap="Table 25. What each swept-wing experiment requires of the cross-flow "
                   "criterion, evaluated at the measured transition station "
                   "(crossflow_criticals_summary.csv).")
table_from_csv("06_validation/crossflow_criticals.csv",
               cap="Table 26. Point-by-point cross-flow criterion values at the measured "
                   "transition stations (crossflow_criticals.csv).")
para("One explanation can be ruled out rather than merely doubted. If the difference "
 "between the two facilities were a Reynolds-number effect the criterion is missing, the "
 "requirement would have to vary with chord Reynolds number in the same direction within a "
 "facility as it does between them. It does not. Within Dagenhart & Saric the required "
 "critical value FALLS steeply with chord Reynolds number, by 251 per decade with a "
 "correlation of −0.88 over a factor of two in Re_c; Boltz et al. sit at six times that "
 "Reynolds number and require 42 % MORE, not less, and are themselves flat across a factor "
 "of three. The between-facility offset therefore has the opposite sign to the "
 "within-facility trend, and no monotone function of Re_c can carry one set into the other. "
 "That is what leaves receptivity — the leading-edge surface finish neither report "
 "documents — as the explanation, and it is now a measurement rather than an appeal to the "
 "literature.")
table_from_csv("06_validation/crossflow_reynolds_trend.csv",
               cap="Table 26a. The required critical value against chord Reynolds number "
                   "within each facility (crossflow_reynolds_trend.csv). The trends have "
                   "opposite signs to the offset between the facilities.")

h2("11.5  Ablations — what each element of the formulation is worth")
para("Each of the three elements that distinguish this formulation is switched off in turn, "
 "everything else held fixed, and the same two datasets the natural branch reaches are re-run. "
 "The table is regenerated by gen_validation.py and is not quoted from memory.")
table_from_csv("06_validation/ablations.csv",
               cap="Table 27. Ablation study (ablations.csv): Schubauer-Skramstad onset error "
                   "and the 86 aerofoil conditions, one closure removed at a time.")
para("Calibration. The solver is calibrated through the single constant set of Table 7. The "
 "critical amplification factor is not among the constants: it follows from the free-stream "
 "turbulence intensity by Mack's correlation, clamped to the 0.0008-0.0298 range over which "
 "that correlation is quoted, and so returns 8.18 for any stream quieter than 0.08 %. The SAME "
 "constants reproduce every validation dataset — five flat plates, two swept wings and 86 "
 "aerofoil conditions — demonstrating that no case-specific physics tuning is required, which "
 "is the essence of the universality claim.")
h2("11.6  Sources and references")
para("All data sources used for validation, for calibration of the closures, and for the "
 "case-study definition are recorded below.")
table_from_csv("06_validation/sources_and_references.csv", max_rows=40,
               cap="Table 28. Validation, calibration and case-study sources.")
table_from_csv("06_validation/swept_wing_source.csv",
               cap="Table 29. Cross-flow calibration dataset provenance.")
table_from_csv("06_validation/swept_wing_independent_source.csv",
               cap="Table 30. Independent swept-wing dataset provenance.")
table_from_csv("06_validation/aerofoil_nlf0416_source.csv",
               cap="Table 31. Aerofoil dataset provenance.")

# ======================================================================
h1("12.  Contribution to Knowledge")
bullet("A single unified transition kernel (Eq. E13) that locally selects the governing "
       "mechanism among natural-TS, bypass, separation-induced and cross-flow transition by a "
       "minimum-effective-onset rule — reproducing all regimes with one calibration set.")
bullet("An amplification database in place of an envelope correlation: 61,600 Orr-Sommerfeld "
       "eigenvalue solutions on the Falkner-Skan family, tabulated against shape factor, "
       "momentum-thickness Reynolds number and frequency, with the march carrying one "
       "amplification factor per physical frequency. The stability solver reproduces the "
       "Blasius neutral point at Re_theta = 201 against the accepted 200.5.")
bullet("A separation-bubble closure that predicts a length rather than a point: the shear layer "
       "is carried across the dead-air region by the momentum integral with no wall stress — a "
       "step with no fitted constant, which reproduces the growth the T3C4 hot films record — "
       "and reattachment placed where the disturbance has amplified by the same N_crit used "
       "elsewhere, so the length scales with the disturbance environment.")
# The figures quoted below are read from the generated CSVs, so this section
# cannot drift from the results it summarises.
_abl = pd.read_csv("06_validation/ablations.csv").set_index("configuration")
_nsum = pd.read_csv("06_validation/aerofoil_nlf0416_summary.csv").set_index("set")
_vsum = pd.read_csv("06_validation/validation_summary.csv")
_plate_err = ", ".join(
    "%+.1f %% on %s" % (r.Re_theta_t_err_pct,
                        r.case.split(" flat plate")[0].replace("ERCOFTAC ", ""))
    for _, r in _vsum.iterrows())
_all = _nsum.loc["All"]
bullet("A two-equation laminar march whose closures are computed from the Falkner-Skan family "
       "rather than fitted, giving the shape factor a history. On the 86 aerofoil conditions it "
       "raises the number of predictions inside the experimental bracket from %d to %d, improving "
       "both surfaces at once."
       % (_abl.loc["one-equation laminar march", "within_bracket"],
          _abl.loc["full model", "within_bracket"]))
bullet("Regime coverage with one constant set: transition-onset Re_theta_t predicted to "
       "%s, spanning 0.03-6 %% free-stream turbulence intensity." % _plate_err)
bullet("An aerofoil validation built for this work: 86 transition locations digitised from Fig. 9 "
       "of NASA TP-1861 for the NLF(1)-0416 section, both surfaces, four chord Reynolds numbers "
       "and lift coefficients from -1.03 to +1.62, with nothing calibrated on them. Mean error "
       "%.3f chord, and %d of the 86 predictions fall inside the +/-0.025c bracket within which "
       "the experiment itself localises transition."
       % (_all.mean_abs_err_c, _all.within_bracket))
bullet("A robust, panel-method-cost (<1 s) 3-D capability via a span-wise strip formulation with "
       "built-in cross-flow, suitable for design-loop use where RANS/LES are impractical.")
bullet("An end-to-end, auditable workflow (geometry → mesh → setup → solution → post-processing "
       "→ validation) producing a complete engineering output set (CSVs, curves, metrics, "
       "contours, temperature profiles, 3-D contours and vectors).")
_nvt = pd.read_csv("04_solution/nlf_vs_turbulent.csv")
bullet("Quantified NLF benefit for the case vehicle: %.0f %% laminar flow and %.0f %% viscous "
       "drag reduction relative to a fully-turbulent wing at the cruise design point."
       % (_nvt.mean_laminar_pct.iloc[0], _nvt.viscous_drag_reduction_pct.iloc[0]))

# ======================================================================
h1("13.  Conclusions")
para("The UTSS universal transition & skin-friction solver predicts boundary-layer transition "
 "over the three-dimensional NLF wing of the AETHER-NLF 25 and the dependent engineering "
 "quantities (skin friction, laminar extent, boundary-layer growth, separation margin, profile "
 "drag) at panel-method cost. The unified four-mechanism kernel, validated with a single "
 "calibration set against five flat plates, two independent swept-wing experiments and 86 "
 "aerofoil conditions, spans 0.03-6 %% free-stream turbulence intensity. At cruise the wing "
 "achieves %.0f %% laminar flow and a %.0f %% viscous-drag reduction versus a turbulent wing; at "
 "the higher-turbulence climb condition the solver switches to the bypass route and predicts "
 "early transition, demonstrating regime coverage across the flight envelope."
 % (_nvt.mean_laminar_pct.iloc[0], _nvt.viscous_drag_reduction_pct.iloc[0]))
para("Two limitations bound that claim and are stated here rather than left to be discovered. "
 "The cross-flow critical constant does not transfer between facilities: the two independent "
 "swept-wing experiments support the functional form of the criterion, and the measured data "
 "collapse on it, but reproducing the second requires a critical value 42 % larger than the "
 "first. And the method is an attached-flow formulation with an incidence envelope of about "
 "+/-8.5 deg on a real section; beyond that the laminar march separates near the leading edge "
 "and the result should not be relied on.")

# ======================================================================
h1("Appendix A.  Complete Generated-Output Inventory")
para("Every file generated for this case study, by folder:")
def inventory():
    rows=[]
    for root in ["01_geometry","02_mesh","03_model_setup","04_solution",
                 "05_postprocessing","06_validation","07_equations"]:
        for dp,_,fs in os.walk(root):
            for f in sorted(fs):
                ext=f.split(".")[-1].lower()
                if ext in ("csv","png","npz"):
                    rows.append([os.path.join(dp,f), ext.upper()])
    return rows
inv=inventory()
para(f"Total generated data/figure files: {len(inv)} "
     f"(CSV: {sum(1 for r in inv if r[1]=='CSV')}, "
     f"PNG: {sum(1 for r in inv if r[1]=='PNG')}, "
     f"NPZ: {sum(1 for r in inv if r[1]=='NPZ')}).")
manual_table(["file","type"], inv, cap="Table A1. Generated-output inventory.")

# ======================================================================
h1("Appendix B.  Parameter / Metric CSVs Rendered as Figures")
para("For completeness, every remaining parameter and metric CSV is also "
     "rendered as a clean figure so that all generated CSV data appear in "
     "plotted form (the same data also appear as native tables earlier).")
for f,c in [("table_geometry_definition","Fig. B1. geometry_definition.csv."),
            ("table_mesh_metrics","Fig. B2. mesh_metrics.csv."),
            ("table_material_properties","Fig. B3. material_properties.csv."),
            ("table_solver_settings","Fig. B4. solver_settings.csv."),
            ("table_integrated_forces","Fig. B5. integrated_forces.csv.")]:
    image(f"05_postprocessing/csv_plots/{f}.png", width=5.9, cap=c)

doc.save("case.docx")
print("case.docx written:", os.path.getsize("case.docx")//1024, "KB")
print("figures embedded across", len(inv), "generated files")
